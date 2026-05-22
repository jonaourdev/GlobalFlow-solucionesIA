from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Iterable, Optional

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document

from app.config import resolve_documentation_dir, resolve_data_file, settings


def normalize_column_name(value: str) -> str:
    value = str(value).strip().lower()
    value = "".join(
        c for c in unicodedata.normalize("NFD", value) if unicodedata.category(c) != "Mn"
    )
    value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
    return value


def normalize_dataframe_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [normalize_column_name(c) for c in df.columns]
    return df


def find_documentation_file(filename: str) -> Optional[Path]:
    documentation_dir = resolve_documentation_dir()
    direct = documentation_dir / filename
    if direct.exists():
        return direct.resolve()

    # Búsqueda flexible por si el nombre cambia levemente.
    stem = Path(filename).stem.lower()
    suffix = Path(filename).suffix.lower()
    for path in documentation_dir.glob(f"*{suffix}"):
        if stem in path.stem.lower() or path.stem.lower() in stem:
            return path.resolve()
    return None


def read_excel_flexible(path: Path) -> pd.DataFrame:
    """Lee un Excel aunque la información esté en la primera hoja disponible."""
    sheets = pd.read_excel(path, sheet_name=None, engine="openpyxl")
    for _, df in sheets.items():
        if not df.empty:
            df = df.dropna(how="all")
            df = df.dropna(axis=1, how="all")
            return normalize_dataframe_columns(df)
    return pd.DataFrame()


def get_tariff_dataframe() -> pd.DataFrame:
    """Prioriza la base arancelaria del repo; mantiene fallback a data/aranceles.csv."""
    excel_path = find_documentation_file(settings.tariff_excel_filename)
    if excel_path:
        return read_excel_flexible(excel_path)

    csv_path = resolve_data_file("aranceles.csv")
    if csv_path.exists():
        return normalize_dataframe_columns(pd.read_csv(csv_path))

    return pd.DataFrame()


def get_historical_dataframe() -> pd.DataFrame:
    """Prioriza las facturas históricas del repo; mantiene fallback a data/historicos.csv."""
    excel_path = find_documentation_file(settings.historical_excel_filename)
    if excel_path:
        return read_excel_flexible(excel_path)

    csv_path = resolve_data_file("historicos.csv")
    if csv_path.exists():
        return normalize_dataframe_columns(pd.read_csv(csv_path))

    return pd.DataFrame()


def docx_to_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def row_to_text(row: pd.Series) -> str:
    values = []
    for column, value in row.items():
        if pd.notna(value):
            values.append(f"{column}: {value}")
    return "\n".join(values)


def load_repository_documents() -> list[Document]:
    """Carga documentos reales del repositorio GlobalFlow-solucionesIA/documentation."""
    docs: list[Document] = []
    documentation_dir = resolve_documentation_dir()

    # Manual normativo principal en Word.
    normative_path = find_documentation_file(settings.normative_docx_filename)
    if normative_path:
        docs.append(
            Document(
                page_content=docx_to_text(normative_path),
                metadata={
                    "tipo": "manual_normativo",
                    "source": str(normative_path),
                    "filename": normative_path.name,
                },
            )
        )

    # Otros .docx dentro de documentation, por ejemplo Documentación Caso GlobalFlow Logistics.docx.
    if documentation_dir.exists():
        for path in documentation_dir.glob("*.docx"):
            if normative_path and path.resolve() == normative_path.resolve():
                continue
            docs.append(
                Document(
                    page_content=docx_to_text(path),
                    metadata={
                        "tipo": "documentacion_caso",
                        "source": str(path),
                        "filename": path.name,
                    },
                )
            )

    # Facturas históricas como documentos recuperables por RAG.
    historical_df = get_historical_dataframe()
    if not historical_df.empty:
        for idx, row in historical_df.iterrows():
            docs.append(
                Document(
                    page_content=row_to_text(row),
                    metadata={
                        "tipo": "historico",
                        "source": settings.historical_excel_filename,
                        "row_index": int(idx),
                    },
                )
            )

    # Base arancelaria también se indexa para búsqueda semántica complementaria.
    tariff_df = get_tariff_dataframe()
    if not tariff_df.empty:
        for idx, row in tariff_df.iterrows():
            docs.append(
                Document(
                    page_content=row_to_text(row),
                    metadata={
                        "tipo": "base_arancelaria",
                        "source": settings.tariff_excel_filename,
                        "row_index": int(idx),
                    },
                )
            )

    return docs
